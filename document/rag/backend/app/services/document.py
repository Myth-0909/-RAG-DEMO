import asyncio
import io
import json
import logging
from typing import AsyncGenerator, Dict, Any, Optional

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.config import settings
from app.models.knowledge import Document, DocumentChunk
from app.services.chunking import chunk_text
from app.services.embedding import embed_texts
from app.services.milvus_service import MilvusService
from app.services.content_analyzer import clean_content
from app.services.llm_analyzer import (
    assess_cleaning_stream,
    analyze_chunking_strategy_stream,
)

logger = logging.getLogger(__name__)

connect_args = {}
if settings.DATABASE_URL.startswith("sqlite"):
    connect_args["check_same_thread"] = False

_engine = create_engine(settings.DATABASE_URL, connect_args=connect_args)
_SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=_engine)


# ── OCR / Table helpers (unchanged) ──────────────────────────────────────────

def _ocr_image(image_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        return text.strip()
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def _table_to_markdown(headers: list, rows: list) -> str:
    if not headers and not rows:
        return ""
    headers = [str(h).strip() if h else "" for h in headers]
    md = "| " + " | ".join(headers) + " |\n"
    md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    for row in rows:
        cells = [str(c).strip() if c else "" for c in row]
        while len(cells) < len(headers):
            cells.append("")
        md += "| " + " | ".join(cells[:len(headers)]) + " |\n"
    return md


def _extract_pdf_tables(file_path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed, skipping PDF table extraction")
        return ""

    table_texts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for j, table in enumerate(tables):
                    if not table or len(table) < 2:
                        continue
                    headers = table[0]
                    rows = table[1:]
                    md = _table_to_markdown(headers, rows)
                    if md:
                        table_texts.append(f"[表格 第{i+1}页 表{j+1}]\n{md}")
    except Exception as e:
        logger.warning(f"PDF table extraction failed: {e}")

    return "\n\n".join(table_texts)


def _extract_pdf_images(file_path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""

    image_texts = []
    try:
        reader = PdfReader(file_path)
        img_idx = 0
        for page_num, page in enumerate(reader.pages):
            for img_key in page.images:
                img_idx += 1
                try:
                    image_data = img_key.data
                    text = _ocr_image(image_data)
                    if text:
                        image_texts.append(f"[图片 第{page_num+1}页 图{img_idx}]\n{text}")
                except Exception as e:
                    logger.warning(f"PDF image {img_idx} extraction failed: {e}")
    except Exception as e:
        logger.warning(f"PDF image extraction failed: {e}")

    return "\n\n".join(image_texts)


def _extract_docx_tables(file_path: str) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return ""

    table_texts = []
    try:
        doc = DocxDocument(file_path)
        for i, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)
            if len(rows_data) < 2:
                continue
            headers = rows_data[0]
            rows = rows_data[1:]
            md = _table_to_markdown(headers, rows)
            if md:
                table_texts.append(f"[表格 {i+1}]\n{md}")
    except Exception as e:
        logger.warning(f"DOCX table extraction failed: {e}")

    return "\n\n".join(table_texts)


def _extract_docx_images(file_path: str) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return ""

    image_texts = []
    try:
        doc = DocxDocument(file_path)
        img_idx = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_idx += 1
                try:
                    image_data = rel.target_part.blob
                    text = _ocr_image(image_data)
                    if text:
                        image_texts.append(f"[图片 {img_idx}]\n{text}")
                except Exception as e:
                    logger.warning(f"DOCX image {img_idx} OCR failed: {e}")
    except Exception as e:
        logger.warning(f"DOCX image extraction failed: {e}")

    return "\n\n".join(image_texts)


def extract_text(file_path: str, file_type: str) -> str:
    """Extract text, tables, and images from a document."""
    parts = []

    if file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = [page.extract_text() or "" for page in reader.pages]
        text_content = "\n\n".join(pages)
        if text_content.strip():
            parts.append(text_content)

        table_content = _extract_pdf_tables(file_path)
        if table_content:
            parts.append("\n\n## 文档表格\n\n" + table_content)

        image_content = _extract_pdf_images(file_path)
        if image_content:
            parts.append("\n\n## 文档图片内容\n\n" + image_content)

    elif file_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text_content = "\n\n".join(
            [para.text for para in doc.paragraphs if para.text.strip()]
        )
        if text_content.strip():
            parts.append(text_content)

        table_content = _extract_docx_tables(file_path)
        if table_content:
            parts.append("\n\n## 文档表格\n\n" + table_content)

        image_content = _extract_docx_images(file_path)
        if image_content:
            parts.append("\n\n## 文档图片内容\n\n" + image_content)

    elif file_type in ("txt", "md"):
        with open(file_path, "r", encoding="utf-8") as f:
            parts.append(f.read())

    else:
        raise ValueError(f"不支持的文件类型: {file_type}")

    return "\n".join(parts)


# ── SSE Event Helpers ────────────────────────────────────────────────────────

def _event(step: str, status: str = "thinking", data: Optional[Dict] = None) -> Dict[str, Any]:
    """Create a standardized SSE event."""
    return {"step": step, "status": status, "data": data or {}}


# ── Main Processing Pipeline (SSE-capable) ───────────────────────────────────

async def process_document_stream(
    doc_id: int,
    file_path: str,
    metadata: dict = None,
) -> AsyncGenerator[Dict[str, Any], None]:
    """
    Process an uploaded document with LLM-driven analysis.

    Yields SSE events showing the model's thinking process:
      1. extract  - Text extraction from file
      2. clean    - LLM assesses whether cleaning is needed
      3. apply    - Apply cleaning if needed
      4. analyze  - LLM analyzes content and outputs chunking plan JSON
      5. chunk    - Execute chunking based on LLM's plan
      6. embed    - Embed chunks and store in Milvus
      7. complete - Done

    Each step emits thinking tokens (via on_thinking callback) so the
    frontend can display the model's reasoning in real-time.
    """
    db = _SessionLocal()
    loop = asyncio.get_running_loop()

    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            yield _event("error", "error", {"message": f"Document {doc_id} not found"})
            return

        doc.status = "processing"
        db.commit()

        # ── Step 1: Extract text ──────────────────────────────────────────
        yield _event("extract", "thinking")

        raw_text = await loop.run_in_executor(None, extract_text, file_path, doc.file_type)

        if not raw_text.strip():
            doc.status = "failed"
            doc.error_message = "文档内容为空"
            db.commit()
            yield _event("extract", "error", {"message": "文档内容为空"})
            return

        yield _event("extract", "done", {
            "chars": len(raw_text),
            "preview": raw_text[:300],
        })

        # ── Step 2: LLM cleaning assessment ───────────────────────────────
        yield _event("clean", "thinking")

        # Bridge sync LLM streaming to async via a queue
        clean_queue: asyncio.Queue = asyncio.Queue()

        def _clean_thinking(token: str):
            loop.call_soon_threadsafe(clean_queue.put_nowait, token)

        def _run_assess():
            try:
                return assess_cleaning_stream(raw_text, doc.file_type, _clean_thinking)
            except Exception as e:
                logger.error(f"Cleaning assessment error: {e}")
                return {
                    "needs_cleaning": False,
                    "issues_found": [f"LLM 分析异常: {e}"],
                    "cleaning_actions": [],
                    "severity": "none",
                    "reasoning": "LLM 调用失败，跳过清洗评估",
                }
            finally:
                loop.call_soon_threadsafe(clean_queue.put_nowait, None)  # sentinel

        assess_task = loop.run_in_executor(None, _run_assess)

        # Yield thinking tokens as they arrive
        while True:
            token = await clean_queue.get()
            if token is None:
                break
            yield _event("clean", "thinking", {"token": token})

        cleaning_assessment = await assess_task

        yield _event("clean", "done", {
            "assessment": cleaning_assessment,
        })

        # ── Step 3: Apply cleaning if needed ──────────────────────────────
        cleaned_text = raw_text
        cleaning_report = {"chars_removed": 0, "skipped": True}

        if cleaning_assessment.get("needs_cleaning", False):
            yield _event("apply", "thinking", {
                "message": "正在执行数据清洗...",
                "actions": cleaning_assessment.get("cleaning_actions", []),
            })

            cleaned_text, cleaning_report = await loop.run_in_executor(
                None, clean_content, raw_text, doc.file_type,
            )

            if not cleaned_text.strip():
                doc.status = "failed"
                doc.error_message = "清洗后文档内容为空"
                db.commit()
                yield _event("apply", "error", {"message": "清洗后文档内容为空"})
                return

            yield _event("apply", "done", {
                "report": cleaning_report,
                "chars_before": len(raw_text),
                "chars_after": len(cleaned_text),
            })
        else:
            yield _event("apply", "done", {
                "message": "无需清洗，跳过",
                "report": cleaning_report,
            })

        # ── Step 4: LLM chunking analysis ─────────────────────────────────
        yield _event("analyze", "thinking")

        chunk_queue: asyncio.Queue = asyncio.Queue()

        def _chunk_thinking(token: str):
            loop.call_soon_threadsafe(chunk_queue.put_nowait, token)

        def _run_analyze():
            try:
                return analyze_chunking_strategy_stream(
                    cleaned_text, doc.file_type, _chunk_thinking,
                )
            except Exception as e:
                logger.error(f"Chunking analysis error: {e}")
                return {
                    "strategy": "recursive",
                    "chunk_size": 500,
                    "chunk_overlap": 100,
                    "analysis": {
                        "document_type": "未分析",
                        "recommended_reasoning": f"LLM 调用失败（{e}），使用默认递归分块",
                    },
                }
            finally:
                loop.call_soon_threadsafe(chunk_queue.put_nowait, None)

        analyze_task = loop.run_in_executor(None, _run_analyze)

        while True:
            token = await chunk_queue.get()
            if token is None:
                break
            yield _event("analyze", "thinking", {"token": token})

        chunking_plan = await analyze_task

        yield _event("analyze", "done", {
            "plan": chunking_plan,
        })

        # ── Step 5: Execute chunking based on LLM plan ────────────────────
        yield _event("chunk", "thinking", {
            "message": f"正在使用 {chunking_plan['strategy']} 策略分块...",
        })

        strategy = chunking_plan["strategy"]
        chunk_size = chunking_plan.get("chunk_size", 500)
        chunk_overlap = chunking_plan.get("chunk_overlap", 100)

        chunks = await loop.run_in_executor(
            None,
            lambda: chunk_text(
                cleaned_text,
                strategy=strategy,
                metadata=None,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            ),
        )

        if not chunks:
            doc.status = "failed"
            doc.error_message = "分块结果为空"
            db.commit()
            yield _event("chunk", "error", {"message": "分块结果为空"})
            return

        yield _event("chunk", "done", {
            "count": len(chunks),
            "strategy": strategy,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
        })

        # ── Step 6: Embed and store in Milvus ─────────────────────────────
        yield _event("embed", "thinking", {
            "message": f"正在向量化 {len(chunks)} 个分块并存入 Milvus...",
        })

        collection_name = f"kb_{doc.knowledge_base_id}"
        milvus = MilvusService()
        milvus.create_collection(collection_name)

        batch_size = 50
        all_milvus_ids = []

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            texts = [c.text for c in batch]

            embeddings = await loop.run_in_executor(None, embed_texts, texts)

            milvus_ids = milvus.insert(
                collection_name=collection_name,
                embeddings=embeddings,
                chunk_texts=texts,
                parent_texts=[c.parent_text for c in batch],
                metadata_list=[c.metadata for c in batch],
                document_ids=[doc_id] * len(batch),
                chunk_indices=[c.chunk_index for c in batch],
            )
            all_milvus_ids.extend(milvus_ids)

            for j, chunk in enumerate(batch):
                db_chunk = DocumentChunk(
                    document_id=doc_id,
                    chunk_index=chunk.chunk_index,
                    chunk_text=chunk.text,
                    parent_text=chunk.parent_text,
                    metadata_json=chunk.metadata,
                    milvus_id=str(all_milvus_ids[i + j]) if i + j < len(all_milvus_ids) else None,
                    token_count=len(chunk.text),
                )
                db.add(db_chunk)

        yield _event("embed", "done", {
            "chunks_embedded": len(chunks),
        })

        # ── Step 7: Update document and finish ────────────────────────────
        strategy_labels = {
            "fixed": "固定大小分块",
            "recursive": "递归字符分块",
            "parent_child": "父子分块",
            "semantic": "语义分块",
            "hybrid": "混合分块",
        }

        doc.status = "completed"
        doc.chunk_count = len(chunks)
        doc.metadata_json = {
            **(doc.metadata_json or {}),
            "cleaning": {
                **cleaning_report,
                "llm_assessment": cleaning_assessment,
            },
            "llm_analysis": chunking_plan,
            "content_analysis": {
                "total_chars": len(cleaned_text),
            },
            "strategy": {
                "selected": strategy,
                "label": strategy_labels.get(strategy, strategy),
                "reasoning": chunking_plan.get("analysis", {}).get(
                    "recommended_reasoning", "LLM 分析推荐"
                ),
            },
        }
        db.commit()

        yield _event("complete", "done", {
            "chunk_count": len(chunks),
            "strategy": strategy,
            "strategy_label": strategy_labels.get(strategy, strategy),
        })

        logger.info(
            f"Document {doc_id} processed: {len(chunks)} chunks "
            f"(strategy: {strategy}, LLM-driven)"
        )

    except Exception as e:
        logger.exception(f"Error processing document {doc_id}")
        try:
            doc = db.query(Document).filter(Document.id == doc_id).first()
            if doc:
                doc.status = "failed"
                doc.error_message = str(e)[:500]
                db.commit()
        except Exception:
            pass
        yield _event("error", "error", {"message": str(e)[:500]})
    finally:
        db.close()


# ── Legacy sync wrapper (for background tasks without SSE) ───────────────────

def process_document(doc_id: int, file_path: str, metadata: dict = None):
    """Sync wrapper for backward compatibility (non-SSE processing)."""
    async def _run():
        async for _event in process_document_stream(doc_id, file_path, metadata):
            pass

    asyncio.run(_run())
